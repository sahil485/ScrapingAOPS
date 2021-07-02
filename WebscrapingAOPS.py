import requests
from urllib.parse import urlparse, urljoin
import colorama
import urllib.request
import webbrowser
import docx
from bs4 import BeautifulSoup
from selenium import webdriver

splitlist = []

def splitlinks(firstlist, finallist):	#function to split links to find actual url of problems
	for i in range(len(firstlist)):
		parts = firstlist[i-1]
		parts = str(parts).split('"')
		finallist.append(parts)


problist =[]

genlink = "https://artofproblemsolving.com"

parameter = "/wiki/index.php"

geo_link = 'https://artofproblemsolving.com/wiki/index.php?title=Category:Introductory_Geometry_Problems&pagefrom=2010+AMC+12A+Problems%2FProblem+17#mw-pages'
cp_link = 'https://artofproblemsolving.com/wiki/index.php/Category:Introductory_Combinatorics_Problems'
alg_link = 'https://artofproblemsolving.com/wiki/index.php?title=Category:Introductory_Algebra_Problems&pagefrom=2001+AIME+II+Problems%2FProblem+1#mw-pages'
nt_link = 'https://artofproblemsolving.com/wiki/index.php/Category:Intermediate_Number_Theory_Problems'

r = requests.get(nt_link).text #gets link from page
soup1 = BeautifulSoup(r, 'lxml') #parses it using lxml so that it is readable

file = open('text.html', encoding='utf8')
soup1 = BeautifulSoup(file, 'lxml')

problems = soup1.find('div', class_='mw-category-group') # finds the specified class of html where all of the problem links are

count = 0

for a_tag in problems.findAll("a"): #finds all "<a href" instances in html and adds the line to the list
	problist.append(a_tag)
	count = count + 1

splitlinks(problist, splitlist)

finallinks = []

for i in range(0, len(splitlist)):	#since there are elements of the split list that are irrelevant to getting the problem, this for loop is to append the actual link query to a list to add to the general URL
	finallinks.append(genlink + splitlist[i][1])

problinks = []

count = 1
 #find amounts of links

driver = webdriver.Chrome(executable_path='C:\webdrivers\chromedriver')

path = 'C:\\Screenshots\\NT\\'


for i in range(1, len(finallinks)+1):

	driver.get(finallinks[i-1])
	driver.save_screenshot(path +str(i+62)+ '.png')

problemsdoc = docx.Document() #saves links of problems with their screenshots on a word document
problemsdoc.add_paragraph("Problem links:")
for i in range(0,len(finallinks)):
	problemsdoc.add_paragraph(finallinks[i])
problemsdoc.save("C:/WordCode/Links")

